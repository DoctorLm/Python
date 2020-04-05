#Python-docx:创建更新Microsoft Word文件的第三方库
#提供创建更新.doc.docx等文件的计算功能

from docx import Document
document = Document()
document.add_heading('Document Title',0)
p = document.add_paragraph('A plain paragraph having some')
document.add_page_break()
document.save('demo.docx')

